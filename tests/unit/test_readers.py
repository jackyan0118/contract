"""Word模板读取器单元测试."""

import pytest
import tempfile
from pathlib import Path
from unittest.mock import Mock, MagicMock, patch, mock_open
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.readers.word_template_reader import WordTemplateReader, TemplateMetadata, TablePlaceholder


class TestWordTemplateReader:
    """WordTemplateReader 测试"""

    def test_reader_creation_default(self):
        """测试默认创建"""
        reader = WordTemplateReader()
        assert reader.template_dir == Path("docs/template")

    def test_reader_creation_custom_dir(self):
        """测试自定义目录"""
        reader = WordTemplateReader(template_dir="custom/path")
        assert reader.template_dir == Path("custom/path")

    @patch("src.readers.word_template_reader.Document")
    @patch("src.readers.word_template_reader.Path.exists")
    def test_read_document(self, mock_exists, mock_document_class):
        """测试读取文档"""
        mock_exists.return_value = True

        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.tables = []
        mock_document_class.return_value = mock_doc

        reader = WordTemplateReader()
        metadata = reader.read("test.docx")

        assert isinstance(metadata, TemplateMetadata)
        assert metadata.file == "test.docx"

    @patch("src.readers.word_template_reader.Document")
    @patch("src.readers.word_template_reader.Path.exists")
    def test_read_with_paragraph_placeholders(self, mock_exists, mock_document_class):
        """测试读取带段落占位符的文档"""
        mock_exists.return_value = True

        # Mock 段落
        mock_para1 = MagicMock()
        mock_para1.text = "客户名称: {{客户名称}}"

        mock_para2 = MagicMock()
        mock_para2.text = "无占位符"

        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2]
        mock_doc.tables = []
        mock_document_class.return_value = mock_doc

        reader = WordTemplateReader()
        metadata = reader.read("test.docx")

        assert "客户名称" in metadata.paragraph_placeholders

    @patch("src.readers.word_template_reader.Document")
    @patch("src.readers.word_template_reader.Path.exists")
    def test_read_file_not_found(self, mock_exists, mock_document_class):
        """测试文件不存在"""
        mock_exists.return_value = False

        reader = WordTemplateReader()

        with pytest.raises(FileNotFoundError, match="not found"):
            reader.read("nonexistent.docx")


class TestTemplateMetadata:
    """TemplateMetadata 测试"""

    def test_metadata_creation(self):
        """测试创建元数据"""
        metadata = TemplateMetadata(
            id="模板1",
            name="测试模板",
            file="test.docx",
            table=None,
            paragraph_placeholders={"客户名称": "", "金额": ""},
            has_speech=False
        )

        assert metadata.id == "模板1"
        assert metadata.name == "测试模板"
        assert "客户名称" in metadata.paragraph_placeholders

    def test_metadata_defaults(self):
        """测试默认参数"""
        metadata = TemplateMetadata(
            id="模板1",
            name="测试",
            file="test.docx"
        )

        assert metadata.table is None
        assert metadata.paragraph_placeholders == {}
        assert metadata.has_speech is False


class TestTablePlaceholder:
    """TablePlaceholder 测试"""

    def test_table_placeholder_creation(self):
        """测试创建表格占位符"""
        placeholder = TablePlaceholder(
            start_pattern="{{#明细表}}",
            end_pattern="{{/明细表}}",
            columns=[{"name": "产品名称", "type": "string"}]
        )

        assert placeholder.start_pattern == "{{#明细表}}"
        assert placeholder.end_pattern == "{{/明细表}}"
        assert len(placeholder.columns) == 1

    def test_table_placeholder_defaults(self):
        """测试默认参数"""
        placeholder = TablePlaceholder(
            start_pattern="{{#表}}",
            end_pattern="{{/表}}",
            columns=[]
        )

        # 验证基本属性存在
        assert placeholder.start_pattern == "{{#表}}"
        assert placeholder.end_pattern == "{{/表}}"


class TestWordTemplateReaderExtended:
    """WordTemplateReader 扩展测试"""

    @patch("src.readers.word_template_reader.Document")
    @patch("src.readers.word_template_reader.Path.exists")
    def test_read_with_table_placeholders(self, mock_exists, mock_document_class):
        """测试读取带表格占位符的文档"""
        mock_exists.return_value = True

        # Mock 表格
        mock_table = MagicMock()
        mock_row = MagicMock()
        mock_cell = MagicMock()
        mock_cell.text = "{{序号}}{{产品名称}}"
        mock_row.cells = [mock_cell]
        mock_table.rows = [mock_row]

        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.tables = [mock_table]
        mock_document_class.return_value = mock_doc

        reader = WordTemplateReader()
        metadata = reader.read("test.docx")

        # 验证元数据返回
        assert metadata.file == "test.docx"

    @patch("src.readers.word_template_reader.Document")
    @patch("src.readers.word_template_reader.Path.exists")
    def test_read_with_speech_markers(self, mock_exists, mock_document_class):
        """测试读取带话术标记的文档"""
        mock_exists.return_value = True

        # Mock 段落 - 包含话术标记
        mock_para = MagicMock()
        mock_para.text = "{{话术}}"

        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para]
        mock_doc.tables = []
        mock_document_class.return_value = mock_doc

        reader = WordTemplateReader()
        metadata = reader.read("test.docx")

        assert metadata.has_speech is True

    @patch("src.readers.word_template_reader.Document")
    @patch("src.readers.word_template_reader.Path.exists")
    def test_read_empty_document(self, mock_exists, mock_document_class):
        """测试读取空文档"""
        mock_exists.return_value = True

        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.tables = []
        mock_document_class.return_value = mock_doc

        reader = WordTemplateReader()
        metadata = reader.read("empty.docx")

        assert metadata.file == "empty.docx"
        assert metadata.paragraph_placeholders == {}

    @patch("src.readers.word_template_reader.Document")
    @patch("src.readers.word_template_reader.Path.exists")
    def test_extract_multiple_placeholders(self, mock_exists, mock_document_class):
        """测试提取多个占位符"""
        mock_exists.return_value = True

        # Mock 多个段落
        mock_para1 = MagicMock()
        mock_para1.text = "{{客户名称}}"
        mock_para2 = MagicMock()
        mock_para2.text = "{{金额}}"
        mock_para3 = MagicMock()
        mock_para3.text = "普通文本"

        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]
        mock_doc.tables = []
        mock_document_class.return_value = mock_doc

        reader = WordTemplateReader()
        metadata = reader.read("test.docx")

        assert len(metadata.paragraph_placeholders) == 2
