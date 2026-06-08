"""表格行扩展器 - 动态扩展表格行."""

import logging
from typing import List, Dict, Any
from copy import deepcopy

from docx import Document
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as _qn

from src.fillers.format_preserver import FormatPreserver

logger = logging.getLogger(__name__)


class RowExpander:
    """表格行扩展器"""

    def __init__(self):
        self.format_preserver = FormatPreserver()

    @staticmethod
    def _row_cells(row, table: Table) -> list:
        """快速获取 row 的 _Cell 列表（避免 Row.cells 触发 vMerge/gridSpan 的 XPath 查询）"""
        return [_Cell(tc, table) for tc in row._tr.tc_lst]

    @staticmethod
    def _row_cell_at(row, table: Table, idx: int):
        """快速获取 row 的第 idx 个 _Cell"""
        tc_lst = row._tr.tc_lst
        if idx < 0 or idx >= len(tc_lst):
            return None
        return _Cell(tc_lst[idx], table)

    def _set_cell_text(self, cell: Any, text: str) -> None:
        """设置单元格文本，强制使用微软雅黑和小五号字体"""
        from docx.oxml.ns import qn

        font_name = "微软雅黑"
        font_size = Pt(9)

        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""
        if cell.paragraphs:
            para = cell.paragraphs[0]
            if not para.runs:
                para.add_run("")

            run = para.runs[0]
            run.text = str(text)
            run.font.name = font_name
            run.font.size = font_size

            # 直接操作XML设置eastAsia字体
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = run._element.makeelement(qn('w:rPr'), {})
                run._element.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = run._element.makeelement(qn('w:rFonts'), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)

    def expand(
        self,
        table: Table,
        data_list: List[Dict[str, Any]],
        columns: List[Dict[str, str]],
        start_row: int = 1,
        template_row_idx: int = -1,
        replace_placeholder: bool = False,
        has_speech_row: bool = True,
        merge_info: dict = None,
        discount_template: str = None
    ) -> None:
        """扩展表格行

        Args:
            table: 表格对象
            data_list: 数据列表
            columns: 列配置列表
            start_row: 起始行索引（从该行开始插入数据）
            template_row_idx: 模板行索引（作为数据行的格式模板）
            replace_placeholder: 是否替换占位行（True=清空占位行作为第一条数据）
            has_speech_row: 是否有话术行（话术行不参与数据填充，需保留在末尾）
            merge_info: 合并信息字典 {row_idx: (discount, span_count)} 用于垂直合并相同折扣的单元格
            discount_template: 折扣话术模板字符串，如 "肾功和心肌酶...*【{discount}】%"
        """
        logger.info(f"[DEBUG] row_expander.expand called: data_count={len(data_list)}, start_row={start_row}")
        if merge_info is None:
            merge_info = {}

        # 如果未配置 discount_template，则不使用折扣话术功能（不合并单元格显示折扣）
        # 只有明确配置了 discount_template 的模板才使用折扣话术

        if not data_list:
            logger.warning("No data to expand")
            return

        data_count = len(data_list)

        if data_count == 1:
            # 只有1条数据：清空占位行，填充该数据
            if start_row < len(table.rows):
                placeholder_row = table.rows[start_row]
                # 关键优化：用 tc_lst 替代 row.cells（避免 vMerge/gridSpan 的 XPath 查询）
                for cell in self._row_cells(placeholder_row, table):
                    self._set_cell_text(cell, "")
                self._fill_row(placeholder_row, table, data_list[0], columns, 1, False, None, discount_template)
            logger.info(f"Filled 1 row (replaced placeholder)")
        else:
            # 大于1条数据：
            # 1. 找到话术行（包含 {{#话术}} 或 {{话术}}），保存为 Tr 元素
            # 2. 保存话术行内容
            # 3. 删除从 start_row 开始的所有数据行（保留标题行和话术行）
            # 4. 在话术行之前插入 N 行数据
            # 5. 填充数据行
            # 6. 对话术行进行文本替换

            # 1. 找到话术行（保存 Tr 元素而不是索引）
            # 关键优化：用 _tr.tc_lst 替代 row.cells（避免 vMerge/gridSpan 的 XPath 查询）
            speech_row = None
            speech_row_texts = []
            for row in table.rows:
                row_text = "".join(
                    "".join(p.text for p in tc.iter(_qn("w:p")))
                    for tc in row._tr.tc_lst
                )
                if "{{#话术}}" in row_text or "{{话术}}" in row_text:
                    speech_row = row
                    speech_row_texts = [
                        "".join(p.text for p in tc.iter(_qn("w:p")))
                        for tc in row._tr.tc_lst
                    ]
                    break

            # 如果没有找到话术行，使用最后一行
            if not speech_row and len(table.rows) > 1:
                speech_row = table.rows[-1]
                speech_row_texts = [
                    "".join(p.text for p in tc.iter(_qn("w:p")))
                    for tc in speech_row._tr.tc_lst
                ]

            if not speech_row:
                logger.warning("No speech row found")
                return

            # 获取表格的列宽（从 tblGrid 获取，这是 python-docx 打开文档后的标准宽度）
            tbl = table._tbl
            tblGrid = tbl.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblGrid')
            col_widths = []
            if tblGrid is not None:
                for col in tblGrid:
                    width = col.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                    if width:
                        col_widths.append(int(width))
                    else:
                        col_widths.append(0)

            # 2. 保存模板行的边框格式（不包括 tcW，因为 python-docx 会自动调整 tcW 为 tblGrid 的值）
            template_row = table.rows[template_row_idx] if template_row_idx >= 0 else table.rows[start_row]
            # 关键优化：用 _tr.tc_lst 替代 row.cells（避免 vMerge/gridSpan 的 XPath 查询）
            template_cell_formats = []
            for tc in template_row._tr.tc_lst:
                tcPr = tc.find('.//{*}tcPr')
                if tcPr is not None:
                    # 复制整个 tcPr，但会覆盖 tcW 为 tblGrid 的值
                    template_cell_formats.append(deepcopy(tcPr))
                else:
                    template_cell_formats.append(None)

            # 2.5 关键优化：预计算每列要复制的子元素（排除 tcW）
            # 原代码内层循环里 dst_tcPr.find(tag) 是 O(F)，每行每列每子元素都做 = O(N*C*F²)
            # 对新建行 dst_tcPr 仅有 tcW，非 tcW 子元素无重复，预计算后省去 find
            cells_children_to_copy: list[list] = []
            for fmt in template_cell_formats:
                if fmt is None:
                    cells_children_to_copy.append([])
                else:
                    cells_children_to_copy.append(
                        [c for c in fmt if not c.tag.endswith('tcW')]
                    )

            # 3. 删除从 start_row 开始到话术行之前的所有行（不删除话术行本身）
            # 先一次性缓存所有行（避免每次访问 table.rows[idx] 触发 findall，O(N²) → O(N)）
            all_rows_for_remove = list(table.rows)
            rows_to_remove = []
            speech_tr = speech_row._tr
            for idx in range(start_row, len(all_rows_for_remove) - 1):
                row_tr = all_rows_for_remove[idx]._tr
                # 不删除话术行
                if row_tr is speech_tr:
                    continue
                rows_to_remove.append(row_tr)

            for tr in rows_to_remove:
                parent = tr.getparent()
                if parent is not None:
                    parent.remove(tr)

            # 4. 在话术行之前插入 N 行数据
            # 缓存行列表获取列数（避免多次触发 findall）
            _base_rows = list(table.rows)
            # 关键优化：用 _tr.tc_lst 替代 row.cells（避免 vMerge/gridSpan 的 XPath 查询）
            target_cols = len(_base_rows[0]._tr.tc_lst) if _base_rows else 7

            # 先收集所有新行（追加到表末尾，O(N) 总体），最后一次性移动到话术行之前
            # 避免每行都调 tbl.insert(pos, ...)，那是 O(N²)
            new_trs: list = []
            for _ in range(data_count):
                # 创建新行
                new_row = table.add_row()
                # 删除多余的单元格，只保留 target_cols 个
                # 关键优化：用 _tr.tc_lst 替代 new_row.cells
                while len(new_row._tr.tc_lst) > target_cols:
                    tr = new_row._tr
                    tc_list = tr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
                    if len(tc_list) > target_cols:
                        tr.remove(tc_list[-1])

                # 关键优化：用 _tr.tc_lst 替代 new_row.cells（避免 vMerge/gridSpan 的 XPath 查询）
                # 缓存 _Cell 包装结果（_Cell.__init__ 本身很快，不触发额外查询）
                new_cells = [_Cell(tc, table) for tc in new_row._tr.tc_lst]
                # 应用单元格格式
                for cell_idx in range(len(new_cells)):
                    # 获取列宽（从 tblGrid）
                    col_width = col_widths[cell_idx] if cell_idx < len(col_widths) else 0

                    # 复制模板行的格式（边框等）
                    if cell_idx < len(cells_children_to_copy) and cells_children_to_copy[cell_idx]:
                        dst_tc = new_cells[cell_idx]._element
                        dst_tcPr = dst_tc.find('.//{*}tcPr')
                        if dst_tcPr is None:
                            dst_tcPr = OxmlElement('w:tcPr')
                            dst_tc.insert(0, dst_tcPr)

                        # 直接 append 所有预筛选的非 tcW 子元素（无需 find 检查，新建行 dst_tcPr 只有 tcW）
                        for src_child in cells_children_to_copy[cell_idx]:
                            dst_tcPr.append(deepcopy(src_child))

                    # 设置 tcW 为 tblGrid 的列宽
                    if col_width > 0:
                        dst_tc = new_cells[cell_idx]._element
                        dst_tcPr = dst_tc.find('.//{*}tcPr')
                        if dst_tcPr is None:
                            dst_tcPr = OxmlElement('w:tcPr')
                            dst_tc.insert(0, dst_tcPr)
                        # 删除旧的 tcW 并添加新的
                        existing_tcW = dst_tcPr.find('.//{*}tcW')
                        if existing_tcW is not None:
                            dst_tcPr.remove(existing_tcW)
                        new_tcW = OxmlElement('w:tcW')
                        new_tcW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', str(col_width))
                        new_tcW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'dxa')
                        dst_tcPr.append(new_tcW)

                new_trs.append(new_row._tr)

            # 一次性把所有新行移到话术行之前（addprevious 是 O(1) 每次）
            # 倒序遍历以保持原顺序：r0, r1, r2, ..., rN-1
            for tr in reversed(new_trs):
                speech_tr.addprevious(tr)

            # 5. 填充数据行
            # 关键修复：缓存 table.rows，避免每次循环都触发 findall，O(N²) → O(N)
            all_rows = list(table.rows)
            last_valid_idx = len(all_rows) - 1
            for row_idx, data in enumerate(data_list):
                target_idx = start_row + row_idx
                if target_idx < last_valid_idx:
                    row = all_rows[target_idx]
                    # 判断该行是否是合并单元格的主单元格
                    is_merge_start = row_idx in merge_info
                    self._fill_row(row, table, data, columns, row_idx + 1, is_merge_start, merge_info, discount_template)

            # 6. 对话术行进行文本替换（保留原有格式）
            if speech_row_texts:
                # 复用 all_rows 缓存，避免再次 findall
                last_row = all_rows[-1]
                # 关键优化：用 _tr.tc_lst 替代 last_row.cells（避免 vMerge/gridSpan 的 XPath 查询）
                last_row_tcs = last_row._tr.tc_lst
                for cell_idx, text in enumerate(speech_row_texts):
                    if cell_idx < len(last_row_tcs):
                        self._set_cell_text(_Cell(last_row_tcs[cell_idx], table), text)

            # 7. 按折扣率合并单元格（供货价列）
            # merge_info 格式：{start_idx: (template_string, span_count)}
            if merge_info:
                self._merge_vertical_cells_by_discount(table, start_row, merge_info, len(columns))

        logger.info(f"Expanded table with {data_count} rows")

    def _merge_vertical_cells_by_discount(self, table: Table, start_row: int, merge_info: dict, col_count: int) -> None:
        """按折扣率垂直合并单元格

        使用vMerge来合并单元格：
        - 第一行使用 vMerge, val=restart（有文本）
        - 后续行使用 vMerge, val=merge（无文本）

        Args:
            table: 表格对象
            start_row: 起始行索引
            merge_info: 合并信息 {row_idx: (template_str, span_count)}
            col_count: 列数量
        """
        if not merge_info or col_count < 3:
            return

        # 供货价列是最后一列（索引 col_count - 1）
        price_col_idx = col_count - 1

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # 关键优化：缓存 table.rows（关键 O(N²) → O(N)）
        # table.rows[i] 内部每次都触发 findall，O(N)；循环内多次访问 = O(N²)
        all_rows = list(table.rows)
        total_rows = len(all_rows)

        # 遍历合并信息
        for row_idx, info in merge_info.items():
            if isinstance(info, tuple) and len(info) == 2:
                template_str, span_count = info
            else:
                continue
            if span_count <= 1:
                continue

            # 计算实际行索引
            actual_start_row = start_row + row_idx
            actual_end_row = actual_start_row + span_count - 1

            # 检查行索引是否有效（包括起始行和结束行）
            if actual_start_row >= total_rows or actual_end_row >= total_rows:
                logger.warning(f"Merge info index out of range: start={actual_start_row}, end={actual_end_row}, table_rows={total_rows}")
                continue

            # 设置合并区域内每一行的边框
            for i in range(span_count):
                current_row_idx = actual_start_row + i
                # 关键优化：用 _tr.tc_lst 替代 row.cells（避免 vMerge/gridSpan 的 XPath 查询）
                tc = all_rows[current_row_idx]._tr.tc_lst[price_col_idx]
                cell = _Cell(tc, table)
                tcPr = tc.find('.//{*}tcPr')
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)

                # 获取或创建 tcBorders
                tcBorders = tcPr.find('.//{*}tcBorders')
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)

                # 判断是否为合并组的第一行或最后一行
                is_first_row = (i == 0)
                is_last_row = (i == span_count - 1)

                # 设置 top 边框（只有第一行可能有）
                if is_first_row:
                    top = tcBorders.find('.//{*}top')
                    if top is None:
                        top = OxmlElement('w:top')
                        tcBorders.append(top)
                    top.set(qn('w:val'), 'nil')
                else:
                    # 后续行的 top 边框应该是 nil
                    top = tcBorders.find('.//{*}top')
                    if top is not None:
                        top.set(qn('w:val'), 'nil')

                # 设置 bottom 边框（只有最后一行需要显示）
                if is_last_row:
                    bottom = tcBorders.find('.//{*}bottom')
                    if bottom is None:
                        bottom = OxmlElement('w:bottom')
                        tcBorders.append(bottom)
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:color'), 'auto')
                    bottom.set(qn('w:sz'), '4')
                else:
                    # 非最后一行，bottom 边框应该是 nil
                    bottom = tcBorders.find('.//{*}bottom')
                    if bottom is not None:
                        tcBorders.remove(bottom)

            # 获取起始行的单元格，设置合并后的文本（保留原有格式）
            # 关键优化：用 _tr.tc_lst 替代 row.cells（避免 vMerge/gridSpan 的 XPath 查询）
            start_tc = all_rows[actual_start_row]._tr.tc_lst[price_col_idx]
            start_cell = _Cell(start_tc, table)
            cell_text = template_str if template_str else ""
            self._set_cell_text(start_cell, cell_text)

            # 设置起始行的 vMerge 为 restart
            tc_start = start_cell._element
            tcPr_start = tc_start.find('.//{*}tcPr')
            if tcPr_start is None:
                tcPr_start = OxmlElement('w:tcPr')
                tc_start.insert(0, tcPr_start)

            # 检查是否已有vMerge
            existing_vMerge = tcPr_start.find('.//{*}vMerge')
            if existing_vMerge is not None:
                tcPr_start.remove(existing_vMerge)

            vMerge_start = OxmlElement('w:vMerge')
            vMerge_start.set(qn('w:val'), 'restart')
            tcPr_start.append(vMerge_start)

            # 从第二行开始设置 vMerge 为 merge
            for merge_row_idx in range(actual_start_row + 1, actual_end_row + 1):
                # 关键优化：用 _tr.tc_lst 替代 row.cells（避免 vMerge/gridSpan 的 XPath 查询）
                merge_tc = all_rows[merge_row_idx]._tr.tc_lst[price_col_idx]
                merge_cell = _Cell(merge_tc, table)

                # 获取tc元素
                tc = merge_tc
                tcPr = tc.find('.//{*}tcPr')
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)

                # 检查是否已有vMerge
                existing_vMerge = tcPr.find('.//{*}vMerge')
                if existing_vMerge is not None:
                    tcPr.remove(existing_vMerge)

                # 添加vMerge元素，值为merge
                vMerge = OxmlElement('w:vMerge')
                vMerge.set(qn('w:val'), 'merge')
                tcPr.append(vMerge)

                # 清空文本（保留格式）
                self._set_cell_text(merge_cell, "")

        logger.info(f"Merged cells for {len(merge_info)} discount groups")

    def _fill_row(
        self,
        row,
        table: Table,
        data: Dict[str, Any],
        columns: List[Dict[str, str]],
        row_number: int,
        is_merge_start: bool = False,
        merge_info: dict = None,
        discount_template: str = None
    ) -> None:
        """填充行数据

        Args:
            row: 表格行
            data: 数据字典
            columns: 列配置
            row_number: 行号
            is_merge_start: 是否是合并单元格的主单元格
            merge_info: 合并信息字典
            discount_template: 折扣话术模板字符串
        """
        if merge_info is None:
            merge_info = {}

        # 关键优化：用 _tr.tc_lst 替代 row.cells（避免 vMerge/gridSpan 的 XPath 查询）
        # 缓存 _Cell 包装结果（_Cell.__init__ 本身很快）
        row_cells = [_Cell(tc, table) for tc in row._tr.tc_lst]
        last_col_idx = len(columns) - 1
        for col_idx, col_config in enumerate(columns):
            if col_idx >= len(row_cells):
                break

            cell = row_cells[col_idx]
            # 判断是否是供货价列（最后一列）
            is_price_col = col_idx == last_col_idx

            if is_price_col:
                # 供货价列处理逻辑
                if is_merge_start and merge_info:
                    # 获取该行的合并信息
                    merge_data = merge_info.get(row_number - 1)  # row_number是1-based
                    if merge_data:
                        template_or_discount, span_count = merge_data
                        # 判断是模板字符串还是折扣值
                        if '{discount}' in template_or_discount:
                            # 旧格式：折扣值，需要格式化
                            discount = data.get('_discount')
                            if discount:
                                value = template_or_discount.format(discount=discount)
                            else:
                                value = ""
                        else:
                            # 新格式：模板字符串直接使用
                            value = template_or_discount
                    else:
                        value = self._get_value(data, col_config, row_number)
                else:
                    # 没有折扣话术时，正常显示供货价
                    value = self._get_value(data, col_config, row_number)
            else:
                value = self._get_value(data, col_config, row_number)

            # 更新单元格（保留原有格式）
            self._set_cell_text(cell, str(value))

    def _merge_cells_by_discount(self, table: Table, start_row: int, merge_info: dict, col_count: int, discount_template: str = None) -> None:
        """按折扣率合并单元格

        Args:
            table: 表格对象
            start_row: 起始行索引
            merge_info: 合并信息 {row_idx: (discount, span_count)}
            col_count: 列数量
            discount_template: 折扣话术模板字符串
        """
        if not merge_info or col_count < 3 or not discount_template:
            return

        # 供货价列是最后一列（索引 col_count - 1）
        price_col_idx = col_count - 1

        # 遍历合并信息
        for row_idx, (discount, span_count) in merge_info.items():
            if span_count <= 1:
                continue

            # 计算实际行索引
            actual_start_row = start_row + row_idx

            # 检查行索引是否有效
            if actual_start_row >= len(table.rows):
                continue

            # 获取起始行和结束行
            start_tr = table.rows[actual_start_row]._tr
            end_idx = actual_start_row + span_count - 1

            if end_idx >= len(table.rows):
                end_idx = len(table.rows) - 1

            # 合并单元格：只合并供货价列（最后一列）
            # 使用 vMerge 来垂直合并单元格
            self._merge_vertical_cells(table, actual_start_row, end_idx, price_col_idx, discount)

        logger.info(f"Merged cells for {len(merge_info)} discount groups")

    def _merge_vertical_cells(self, table: Table, start_row: int, end_row: int, col_idx: int, discount: int) -> None:
        """垂直合并指定列的单元格

        Args:
            table: 表格对象
            start_row: 起始行索引
            end_row: 结束行索引
            col_idx: 列索引
            discount: 折扣率（用于设置合并后的文本）
        """
        if start_row >= end_row:
            return

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn, nsmap

        # 获取起始行和结束行的单元格
        start_cell = table.rows[start_row].cells[col_idx]
        end_cell = table.rows[end_row].cells[col_idx]

        # 设置合并后的文本内容（保留原有格式）
        cell_text = discount_template.format(discount=discount)
        self._set_cell_text(start_cell, cell_text)

        # 从第二行开始，将单元格的 vMerge 设置为 'merge'，表示合并到第一个单元格
        for row_idx in range(start_row + 1, end_row + 1):
            cell = table.rows[row_idx].cells[col_idx]

            # 获取或创建 tcPr 元素
            tc = cell._element
            # 使用命名空间查找
            tcPr = tc.find('.//{%s}tcPr' % nsmap['w'])
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # 添加 vMerge 元素，值为 'merge' 表示合并到上一个单元格
            vMerge = OxmlElement('w:vMerge')
            vMerge.set(qn('w:val'), 'merge')
            tcPr.append(vMerge)

            # 清空单元格的文本（保留格式）
            self._set_cell_text(cell, "")

        logger.info(f"Merged rows {start_row}-{end_row} at column {col_idx} with discount {discount}%")

    def _get_value(
        self,
        data: Dict[str, Any],
        col_config: Dict[str, str],
        row_number: int
    ) -> Any:
        """获取单元格值"""
        col_type = col_config.get("type", "text")
        field = col_config.get("field", "")

        # 自动编号：如果有_group_index（折扣分组），使用组内序号
        if col_type == "auto_number":
            if '_group_index' in data:
                return data['_group_index']
            return row_number

        # 获取字段值
        value = data.get(field, "")

        if value is None:
            return ""

        # 转换
        transform = col_config.get("transform")
        if transform == "substring":
            start = col_config.get("params", {}).get("start", 0)
            length = col_config.get("params", {}).get("length")
            if length:
                return str(value)[start:start + length]
            return str(value)[start:]

        if transform == "currency":
            decimals = col_config.get("params", {}).get("decimals", 2)
            try:
                return f"{float(value):.{decimals}f}"
            except (TypeError, ValueError):
                return value

        return value
