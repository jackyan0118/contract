"""数据填充引擎 - 将数据填充到 Word 模板."""

import logging
import re
import time
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
from decimal import Decimal
from enum import Enum

from docx import Document
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.database.connection import get_connection_pool
from src.database.config import get_database_config

logger = logging.getLogger(__name__)


class Operator(str, Enum):
    """操作符枚举"""
    EQ = "="
    NE = "!="
    CONTAINS = "contains"
    IN = "in"
    NOT_IN = "not_in"
    GT = ">"
    LT = "<"
    GTE = ">="
    LTE = "<="


@dataclass
class FilterCondition:
    """过滤条件"""
    field: str
    operator: str  # =, !=, contains, in, not_in, >, <, >=, <=
    value: Any
    value_type: Optional[str] = None  # 可选：指定值的类型，如 "bm" 表示BM编码需转换为ID


@dataclass
class FilterRule:
    """过滤规则"""
    id: str
    name: str
    conditions: List[FilterCondition] = field(default_factory=list)


@dataclass
class SpeechConfig:
    """话术配置"""
    id: str
    type: str  # conditional, fixed
    mutex_group: Optional[str] = None
    conditions: List[FilterCondition] = field(default_factory=list)
    content: str = ""
    variables: Dict[str, str] = field(default_factory=dict)


@dataclass
class ColumnConfig:
    """列配置"""
    name: str
    source_field: str
    source_table: str = "UF_HTJGKST_DT1"
    type: str = "text"  # text, auto_number
    transform: Optional[str] = None
    params: Dict[str, Any] = field(default_factory=dict)


class DataFiller:
    """数据填充引擎"""

    def __init__(self):
        self.operators = {
            "=": self._op_equals,
            "!=": self._op_not_equals,
            "contains": self._op_contains,
            "in": self._op_in,
            "not_in": self._op_not_in,
            ">": self._op_greater,
            "<": self._op_less,
            ">=": self._op_greater_equal,
            "<=": self._op_less_equal,
        }
        # BM映射表缓存
        self._bm_to_id_cache: Optional[Dict[str, int]] = None
        self._bm_to_id_cache_time: float = 0
        self._bm_to_id_cache_ttl: float = 3600  # 缓存1小时

    def fill_paragraphs(
        self,
        doc: Document,
        placeholders: Dict[str, str],
        data: Dict[str, Any]
    ) -> None:
        """填充段落占位符

        Args:
            doc: Word 文档对象
            placeholders: 占位符字典
            data: 数据字典
        """
        for para in doc.paragraphs:
            text = para.text

            # 替换占位符
            for key, value in placeholders.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in text:
                    # 优先使用传入的值，否则使用占位符配置的值
                    actual_value = data.get(key, value)
                    text = text.replace(placeholder, str(actual_value))

            # 替换变量占位符
            text = self._replace_variables(text, data)

            # 更新段落文本，强制使用微软雅黑和小五号字体
            if text != para.text:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = text
                    para.runs[0].font.name = "微软雅黑"
                    try:
                        para.runs[0].font.eastAsia = "微软雅黑"
                    except Exception:
                        pass
                    para.runs[0].font.size = Pt(9)
                else:
                    new_run = para.add_run(text)
                    new_run.font.name = "微软雅黑"
                    try:
                        new_run.font.eastAsia = "微软雅黑"
                    except Exception:
                        pass
                    new_run.font.size = Pt(9)

    def fill_table(
        self,
        table: Table,
        columns: List[ColumnConfig],
        data_list: List[Dict[str, Any]],
        start_row: int = 1
    ) -> None:
        """填充表格数据

        Args:
            table: Word 表格对象
            columns: 列配置列表
            data_list: 数据行列表
            start_row: 数据起始行索引
        """
        if not data_list:
            logger.warning("No data to fill in table")
            return

        # 确保表格有足够的行
        required_rows = start_row + len(data_list)
        self._ensure_table_rows(table, required_rows)

        # 填充数据行
        for row_idx, data in enumerate(data_list):
            table_row_idx = start_row + row_idx

            if table_row_idx >= len(table.rows):
                break

            row = table.rows[table_row_idx]

            for col_idx, col_config in enumerate(columns):
                if col_idx >= len(row.cells):
                    break

                cell = row.cells[col_idx]
                value = self._get_cell_value(data, col_config)

                # 更新单元格文本，强制使用微软雅黑和小五号字体
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    # 彻底清除所有runs，只保留一个
                    for run in para.runs:
                        run.text = ""
                    # 如果没有runs，添加一个空的
                    if not para.runs:
                        para.add_run("")
                    # 设置新文本
                    para.runs[0].text = str(value)
                    para.runs[0].font.name = "微软雅黑"
                    try:
                        para.runs[0].font.eastAsia = "微软雅黑"
                    except Exception:
                        pass
                    para.runs[0].font.size = Pt(9)

    def filter_data(
        self,
        data_list: List[Dict[str, Any]],
        conditions: List[FilterCondition]
    ) -> List[Dict[str, Any]]:
        """过滤数据

        Args:
            data_list: 原始数据列表
            conditions: 过滤条件列表

        Returns:
            过滤后的数据列表
        """
        if not conditions:
            return data_list

        filtered = []
        for data in data_list:
            if self.match_conditions(data, conditions):
                filtered.append(data)

        logger.info(f"Filtered {len(data_list)} rows to {len(filtered)} rows")
        return filtered

    def match_conditions(
        self,
        data: Dict[str, Any],
        conditions: List[FilterCondition]
    ) -> bool:
        """检查数据是否满足所有条件"""
        for cond in conditions:
            if not self._match_condition(data, cond):
                return False
        return True

    def _match_condition(
        self,
        data: Dict[str, Any],
        condition: FilterCondition
    ) -> bool:
        """检查单个条件"""
        # 获取字段值（考虑字段映射）
        field_value = self._get_field_value(data, condition.field)

        # 处理值类型转换（如 BM 编码转 ID）
        converted_value = self._convert_value(condition.value, condition.value_type, data)

        # 调试日志
        logger.debug(
            f"Matching condition: field={condition.field}, field_value={field_value}, "
            f"value={condition.value}, value_type={condition.value_type}, "
            f"converted_value={converted_value}, operator={condition.operator}"
        )

        # 获取操作符函数
        op_func = self.operators.get(condition.operator)
        if not op_func:
            logger.warning(f"Unknown operator: {condition.operator}")
            return True

        return op_func(field_value, converted_value)

    def _convert_value(self, value: Any, value_type: Optional[str], data: Dict[str, Any]) -> Any:
        """转换值类型

        Args:
            value: 原始值（可以是单个值或列表）
            value_type: 值类型 (如 "bm" 表示BM编码)
            data: 数据字典（用于获取上下文）

        Returns:
            转换后的值
        """
        if value_type == "bm":
            # BM编码转ID：从常量或缓存获取映射
            bm_to_id = self._get_bm_to_id_map()

            # 处理列表值（如 ['11', '41']）
            if isinstance(value, list):
                converted_list = []
                for v in value:
                    bm_value = str(v).strip()
                    if bm_value in bm_to_id:
                        # 转换为字符串，保持与数据库字段类型一致
                        converted_list.append(str(bm_to_id[bm_value]))
                    else:
                        logger.warning(f"BM编码 {bm_value} 未找到对应的ID")
                        converted_list.append(v)
                return converted_list

            # 处理单个值
            bm_value = str(value).strip()
            if bm_value in bm_to_id:
                return str(bm_to_id[bm_value])
            logger.warning(f"BM编码 {bm_value} 未找到对应的ID")
            return value

        return value

    def _get_bm_to_id_map(self) -> Dict[str, int]:
        """获取BM到ID的映射表

        从数据库查询，支持缓存机制避免频繁查询
        """
        # 检查缓存是否有效
        current_time = time.time()
        if self._bm_to_id_cache is not None:
            if current_time - self._bm_to_id_cache_time < self._bm_to_id_cache_ttl:
                return self._bm_to_id_cache

        # 从数据库查询映射
        try:
            pool = get_connection_pool()
            db_config = get_database_config()
            table_name = db_config.get_qualified_table("UF_CPXF")

            with pool.connection() as conn:
                with conn.cursor() as cursor:
                    cursor.execute(
                        f"SELECT BM, ID FROM {table_name} WHERE BM IS NOT NULL AND SFQY = 0"
                    )
                    rows = cursor.fetchall()

            # 构建映射字典
            bm_to_id = {}
            for bm, id_ in rows:
                if bm is not None and id_ is not None:
                    bm_to_id[str(bm).strip()] = int(id_)

            # 更新缓存
            self._bm_to_id_cache = bm_to_id
            self._bm_to_id_cache_time = current_time

            logger.info(f"从数据库加载了 {len(bm_to_id)} 条BM映射记录")
            return bm_to_id

        except Exception as e:
            logger.warning(f"从数据库查询BM映射失败: {e}，使用空映射")
            # 返回空映射，让调用方处理找不到的情况
            return {}

    def _get_field_value(
        self,
        data: Dict[str, Any],
        field: str
    ) -> Any:
        """获取字段值（大小写不敏感查找）"""
        # 直接查找（区分大小写）
        if field in data:
            return data[field]

        # 尝试小写查找（数据库字段可能是大写）
        field_lower = field.lower()
        for key in data:
            if key.lower() == field_lower:
                return data[key]

        return None

    def _op_equals(self, field_value: Any, expected: Any) -> bool:
        """等于操作符"""
        return str(field_value).strip() == str(expected).strip()

    def _op_not_equals(self, field_value: Any, expected: Any) -> bool:
        """不等于操作符"""
        return str(field_value).strip() != str(expected).strip()

    def _op_contains(self, field_value: Any, expected: Any) -> bool:
        """包含操作符"""
        return str(expected) in str(field_value)

    def _op_in(self, field_value: Any, expected: List[Any]) -> bool:
        """在列表中操作符"""
        field_str = str(field_value).strip()
        expected_strs = [str(v).strip() for v in expected]
        return field_str in expected_strs

    def _op_not_in(self, field_value: Any, expected: List[Any]) -> bool:
        """不在列表中操作符"""
        field_str = str(field_value).strip()
        expected_strs = [str(v).strip() for v in expected]
        return field_str not in expected_strs

    def _op_greater(self, field_value: Any, expected: Any) -> bool:
        """大于操作符"""
        try:
            return float(field_value) > float(expected)
        except (TypeError, ValueError):
            return False

    def _op_less(self, field_value: Any, expected: Any) -> bool:
        """小于操作符"""
        try:
            return float(field_value) < float(expected)
        except (TypeError, ValueError):
            return False

    def _op_greater_equal(self, field_value: Any, expected: Any) -> bool:
        """大于等于操作符"""
        try:
            return float(field_value) >= float(expected)
        except (TypeError, ValueError):
            return False

    def _op_less_equal(self, field_value: Any, expected: Any) -> bool:
        """小于等于操作符"""
        try:
            return float(field_value) <= float(expected)
        except (TypeError, ValueError):
            return False

    def _get_cell_value(
        self,
        data: Dict[str, Any],
        col_config: ColumnConfig
    ) -> Any:
        """获取单元格值"""
        field = col_config.source_field

        # 自动编号
        if col_config.type == "auto_number":
            return ""

        # 获取原始值
        value = data.get(field)

        # 应用转换
        if value is None:
            return ""

        if col_config.transform == "substring":
            start = col_config.params.get("start", 0)
            length = col_config.params.get("length")
            if length:
                return str(value)[start:start + length]
            return str(value)[start:]

        if col_config.transform == "currency":
            decimals = col_config.params.get("decimals", 2)
            try:
                return f"{float(value):.{decimals}f}"
            except (TypeError, ValueError):
                return value

        return value

    def _replace_variables(self, text: str, data: Dict[str, Any]) -> str:
        """替换变量占位符"""
        # 匹配 {{变量名}} 格式
        pattern = r'\{\{(\w+)\}\}'
        matches = re.findall(pattern, text)

        for var_name in matches:
            placeholder = f"{{{{{var_name}}}}}"
            value = data.get(var_name, "")
            text = text.replace(placeholder, str(value))

        return text

    def _ensure_table_rows(self, table: Table, required_rows: int) -> None:
        """确保表格有足够的行"""
        current_rows = len(table.rows)

        if current_rows < required_rows:
            # 复制最后一行作为模板
            template_row = table.rows[-1] if table.rows else None

            for _ in range(required_rows - current_rows):
                if template_row:
                    # 复制模板行
                    new_row = table.add_row()
                    for src_cell, dst_cell in zip(template_row.cells, new_row.cells):
                        for src_para, dst_para in zip(src_cell.paragraphs, dst_cell.paragraphs):
                            for src_run, dst_run in zip(src_para.runs, dst_para.runs):
                                new_run = dst_para.add_run(src_run.text)
                                # 强制使用微软雅黑和小五号字体
                                new_run.font.name = "微软雅黑"
                                try:
                                    new_run.font.eastAsia = "微软雅黑"
                                except Exception:
                                    pass
                                new_run.font.size = Pt(9)
                else:
                    table.add_row()
