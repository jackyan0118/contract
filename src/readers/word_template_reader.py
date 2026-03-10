"""Word 模板读取器 - 解析模板文件结构."""

import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


@dataclass
class TablePlaceholder:
    """表格占位符"""
    start_pattern: str = "{{#明细表}}"
    end_pattern: str = "{{/明细表}}"
    columns: List[Dict[str, str]] = field(default_factory=list)


@dataclass
class ParagraphPlaceholder:
    """段落占位符"""
    name: str
    value: str


@dataclass
class SpeechPlaceholder:
    """话术占位符"""
    start_pattern: str = "{{#话术}}"
    end_pattern: str = "{{/话术}}"


@dataclass
class TemplateMetadata:
    """模板元数据"""
    id: str
    name: str
    file: str
    category: str = ""
    table: Optional[TablePlaceholder] = None
    paragraph_placeholders: Dict[str, str] = field(default_factory=dict)
    has_speech: bool = False


class WordTemplateReader:
    """Word 模板读取器 - 解析模板结构"""

    def __init__(self, template_dir: str = "docs/template"):
        self.template_dir = Path(template_dir)

    def read(self, template_file: str) -> TemplateMetadata:
        """读取模板文件并解析结构

        Args:
            template_file: 模板文件路径

        Returns:
            TemplateMetadata: 模板元数据
        """
        template_path = self.template_dir / template_file

        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        doc = Document(template_path)

        # 解析段落占位符
        paragraph_placeholders = self._parse_paragraph_placeholders(doc)

        # 解析表格占位符
        table = self._parse_table_placeholder(doc)

        # 检查话术占位符
        has_speech = self._check_speech_placeholder(doc)

        return TemplateMetadata(
            id="",
            name="",
            file=template_file,
            table=table,
            paragraph_placeholders=paragraph_placeholders,
            has_speech=has_speech
        )

    def _parse_paragraph_placeholders(self, doc: Document) -> Dict[str, str]:
        """解析段落占位符"""
        placeholders = {}

        for para in doc.paragraphs:
            text = para.text
            # 匹配 {{xxx}} 格式的占位符
            matches = re.findall(r'\{\{(\w+)\}\}', text)
            for match in matches:
                placeholders[match] = ""

        return placeholders

    def _parse_table_placeholder(self, doc: Document) -> Optional[TablePlaceholder]:
        """解析表格占位符"""
        # 查找包含 {{#明细表}} 的表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "{{#明细表}}" in para.text:
                            # 找到明细表，解析列
                            columns = self._parse_table_columns(table)
                            return TablePlaceholder(columns=columns)

        return None

    def _parse_table_columns(self, table: Table) -> List[Dict[str, str]]:
        """解析表格列"""
        columns = []

        if not table.rows:
            return columns

        # 第一行为表头
        header_row = table.rows[0]

        for idx, cell in enumerate(header_row.cells):
            text = cell.text.strip()
            # 移除可能的占位符标记
            text = text.replace("{{", "").replace("}}", "")

            # 确定字段名
            field_name = self._infer_field_name(text, idx)
            column_type = self._infer_column_type(text, idx)

            columns.append({
                "name": text,
                "field": field_name,
                "type": column_type
            })

        return columns

    def _infer_field_name(self, header_text: str, col_idx: int) -> str:
        """推断字段名"""
        header_to_field = {
            "序号": "序号",
            "物料编码": "WLDM",
            "SAP代码": "WLDM",
            "品名": "WLMS",
            "产品名称": "WLMS",
            "简称": "WLMS",
            "规格": "GG",
            "规格型号": "GG",
            "包装规格": "GG",
            "零售价": "LSJ",
            "供货价": "GHJY",
            "产品类别": "CPXF",
            "分类": "CPXF",
        }

        for key, value in header_to_field.items():
            if key in header_text:
                return value

        return f"col_{col_idx}"

    def _infer_column_type(self, header_text: str, col_idx: int) -> str:
        """推断列类型"""
        if "序号" in header_text:
            return "auto_number"

        return "text"

    def _check_speech_placeholder(self, doc: Document) -> bool:
        """检查话术占位符"""
        for para in doc.paragraphs:
            if "{{#话术}}" in para.text or "{{话术}}" in para.text:
                return True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "{{话术}}" in para.text:
                            return True

        return False

    def get_table_data_range(self, template_file: str) -> Tuple[int, int]:
        """获取表格数据行范围

        Returns:
            (start_row, end_row): 数据起始和结束行索引
        """
        template_path = self.template_dir / template_file
        doc = Document(template_path)

        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        if "{{#明细表}}" in text:
                            # 找到起始标记，下一行是数据开始
                            return row_idx + 1, len(table.rows) - 1

        return 0, 0
