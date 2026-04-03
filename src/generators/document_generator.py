"""文档生成器 - 根据模板和数据生成 Word 文档."""

import logging
import re
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime

from docx import Document

from src.readers.word_template_reader import WordTemplateReader
from src.fillers.data_filler import DataFiller, FilterCondition
from src.fillers.row_expander import RowExpander
from src.fillers.format_preserver import FormatPreserver
from src.fillers.speech_processor import SpeechProcessor, Speech
from src.fillers.constants import DEFAULT_SPEECH_VARIABLES, HEADER_TO_FIELD
from src.config.template_loader import (
    TemplateLoader,
    TemplateMetadataModel,
    ColumnType,
    SortRuleModel,
)
from src.models.template_rule import TemplateRule

logger = logging.getLogger(__name__)


@dataclass
class GenerationResult:
    """生成结果"""

    success: bool
    file_path: Optional[str] = None
    template_id: Optional[str] = None
    error: Optional[str] = None


class DocumentGenerator:
    """文档生成器"""

    def __init__(
        self,
        template_dir: str = "templates",
        config_dir: str = "config/template_metadata/templates",
    ):
        self.template_dir = Path(template_dir)
        self.config_dir = Path(config_dir)
        self.template_reader = WordTemplateReader(template_dir)
        self.template_loader = TemplateLoader(config_dir)
        self.data_filler = DataFiller()
        self.row_expander = RowExpander()
        self.format_preserver = FormatPreserver()
        self.speech_processor = SpeechProcessor()

    def generate(
        self,
        template: TemplateRule,
        quote_data: Dict[str, Any],
        detail_data_list: List[Dict[str, Any]],
        output_dir: str = "output",
        wybs: str = None,
    ) -> GenerationResult:
        """生成文档

        Args:
            template: 模板规则
            quote_data: 报价单主表数据
            detail_data_list: 报价单明细数据列表
            output_dir: 输出目录
            wybs: 报价单号（用于组织输出目录结构）

        Returns:
            GenerationResult: 生成结果
        """
        # 尝试加载模板配置
        template_config: Optional[TemplateMetadataModel] = None
        try:
            # 优先使用 template.id
            template_config = self.template_loader.load(template.id)
        except FileNotFoundError:
            # 如果找不到，尝试从文件名提取配置名
            # 例如：template.file = "模板2：通用生化产品价格模版.docx" -> config = "模板2.yaml"
            config_name = template.file.replace(".docx", "").split("：")[0]
            try:
                template_config = self.template_loader.load(config_name)
                logger.debug(f"Loaded template config using filename: {config_name}")
            except FileNotFoundError:
                logger.debug(f"No template config found for {config_name}, using default")
            except Exception as e:
                logger.warning(f"Failed to load template config: {e}")
        except Exception as e:
            logger.warning(f"Failed to load template config: {e}")

        try:
            template_path = self.template_dir / template.file

            if not template_path.exists():
                return GenerationResult(
                    success=False,
                    template_id=template.id,
                    error=f"Template file not found: {template.file}",
                )

            # 复制模板到输出目录（按报价单号组织目录结构）
            if wybs:
                output_path = (
                    Path(output_dir)
                    / wybs
                    / f"{template.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
                )
            else:
                output_path = (
                    Path(output_dir)
                    / f"{template.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
                )
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy(template_path, output_path)

            # 打开文档
            doc = Document(output_path)

            # 填充数据（带模板配置）
            self._fill_document(doc, template_config, quote_data, detail_data_list)

            # 保存文档
            doc.save(output_path)

            logger.info(f"Generated document: {output_path}")

            return GenerationResult(
                success=True, file_path=str(output_path), template_id=template.id
            )

        except Exception as e:
            logger.error(f"Failed to generate document: {e}")
            import traceback

            logger.error(traceback.format_exc())
            return GenerationResult(success=False, template_id=template.id, error=str(e))

    def _fill_document(
        self,
        doc: Document,
        template_config: Optional[TemplateMetadataModel],
        quote_data: Dict[str, Any],
        detail_data_list: List[Dict[str, Any]],
    ) -> None:
        """填充文档内容

        Args:
            doc: Word 文档对象
            template_config: 模板配置（可选）
            quote_data: 报价单主表数据
            detail_data_list: 明细数据列表
        """
        # 1. 填充段落占位符
        self._fill_paragraphs(doc, template_config, quote_data)

        # 2. 填充表格（支持 detail_filter）
        self._fill_tables(doc, template_config, detail_data_list, quote_data)

        # 3. 处理话术（完整集成）- 传递明细数据用于话术条件判断
        self._process_speeches(doc, template_config, detail_data_list)

        # 4. 修复表格列宽：将 autofit 改为 fixed，防止 WPS 渲染时列宽随内容变化
        self._fix_table_layout(doc)

    def _fill_paragraphs(
        self, doc: Document, template_config: Optional[TemplateMetadataModel], data: Dict[str, Any]
    ) -> None:
        """填充段落占位符

        Args:
            doc: Word 文档对象
            template_config: 模板配置（可选）
            data: 数据字典
        """
        # 获取段落占位符配置（优先使用模板配置，否则使用默认值）
        if template_config and template_config.paragraph_placeholders:
            placeholders = template_config.paragraph_placeholders
        else:
            placeholders = {
                "标题": data.get("标题", ""),
                "副标题": data.get("副标题", ""),
                "金额单位": data.get("金额单位", "元"),
                "报价单号": data.get("报价单号", ""),
                "客户名称": data.get("客户名称", ""),
            }

        for para in doc.paragraphs:
            text = para.text

            # 替换占位符
            for key, value in placeholders.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in text:
                    text = text.replace(placeholder, str(value))

            # 替换变量占位符（如 {{肝功扣率}}）
            text = self._replace_variables(text, data)

            # 更新段落
            if text != para.text:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = text

    def _fill_tables(
        self,
        doc: Document,
        template_config: Optional[TemplateMetadataModel],
        data_list: List[Dict[str, Any]],
        quote_data: Dict[str, Any] = None,
    ) -> None:
        """填充表格

        Args:
            doc: Word 文档对象
            template_config: 模板配置（可选）
            data_list: 明细数据列表
            quote_data: 报价单主表数据（用于话术处理）
        """
        if quote_data is None:
            quote_data = {}
        for table in doc.tables:
            # 查找明细表和话术行
            is_detail_table = False
            template_row_idx = -1
            speech_row_idx = -1
            speech_row_content = None

            # 第一遍：清除所有占位符并找到关键行
            for row_idx, row in enumerate(table.rows):
                row_text = ""
                for cell in row.cells:
                    for para in cell.paragraphs:
                        row_text += para.text

                # 检查是否是明细表模板行
                if "{{#明细表}}" in row_text:
                    is_detail_table = True
                    template_row_idx = row_idx
                    # 清除标记文本
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            para.text = para.text.replace("{{#明细表}}", "").replace(
                                "{{/明细表}}", ""
                            )

                # 检查是否是话术行
                if "{{#话术}}" in row_text or "{{话术}}" in row_text:
                    speech_row_idx = row_idx
                    # 保存话术行内容
                    speech_row_content = []
                    for cell in row.cells:
                        cell_text = ""
                        for para in cell.paragraphs:
                            cell_text += para.text
                        speech_row_content.append(cell_text)

            # 如果没有找到 {{#明细表}} 占位符，但模板配置了 table.columns，则也视为明细表
            if (
                not is_detail_table
                and template_config
                and template_config.table
                and template_config.table.columns
            ):
                is_detail_table = True
                # 假设第一行是表头，最后一行是话术行或空行
                template_row_idx = 1  # 从第1行开始填充数据
                has_speech_row = speech_row_content is not None

            if not is_detail_table:
                continue

            # 应用 detail_filter 过滤数据
            filtered_data = self._apply_detail_filter(template_config, data_list)

            # 应用排序规则
            sorted_data = self._apply_sort_rules(template_config, filtered_data)

            # 应用去重规则
            deduped_data = self._apply_dedup_rules(template_config, sorted_data)

            # 应用产品匹配
            matched_data = self._apply_product_matching(template_config, deduped_data)

            # 计算折扣分组
            if template_config and (template_config.product_matching or template_config.discount_templates):
                grouped_data, merge_info = self._calculate_discount_and_group(matched_data, template_config)
            else:
                grouped_data = matched_data
                merge_info = {}

            # 解析列配置（优先使用模板配置）
            if template_config and template_config.table:
                columns = self._parse_table_columns_from_config(template_config.table)
            else:
                columns = self._parse_table_columns(table)

            # 填充数据 - 从占位行位置开始（替换占位行）
            # 根据是否有话术行设置 has_speech_row 参数
            has_speech_row = speech_row_content is not None

            # 调试：检查过滤后的数据
            logger.info(
                f"filtered_data count: {len(filtered_data)}, grouped_data count: {len(grouped_data)}, columns[0]: {columns[0] if columns else 'empty'}, template_row_idx: {template_row_idx}"
            )

            # 获取折扣话术模板
            discount_template = template_config.discount_template if template_config else None
            self.row_expander.expand(
                table,
                grouped_data,
                columns,
                template_row_idx,
                template_row_idx,
                True,
                has_speech_row,
                merge_info,
                discount_template,
            )

            # 获取话术内容
            speech_contents = self._get_speech_contents(template_config, data_list)
            full_speech = "\n".join(speech_contents)

            # 替换变量占位符
            for var_name, default_value in DEFAULT_SPEECH_VARIABLES.items():
                full_speech = full_speech.replace(f"{{{{{var_name}}}}}", default_value)

            # 直接找到包含话术占位符的行，替换占位符
            for row_idx, row in enumerate(table.rows):
                row_text = ""
                for cell in row.cells:
                    row_text += cell.text

                if "{{#话术}}" in row_text or "{{话术}}" in row_text:
                    # 找到话术行，替换占位符
                    # 只修改合并单元格的主单元格（Cell 1）
                    # Cell 2-6 是虚拟单元格，由于 gridSpan 合并，它们的文本会被忽略

                    # 修改 Cell 1 - 这是合并单元格的主单元格
                    cell1 = row.cells[1]
                    # 清空并设置新内容
                    cell1.text = full_speech

                    # 修改 Cell 0 - 保留 "其他"，移除占位符
                    cell0 = row.cells[0]
                    cell0_text = (
                        cell0.text.replace("{{#话术}}", "")
                        .replace("{{/话术}}", "")
                        .replace("{{话术}}", "")
                    )
                    cell0.text = cell0_text if cell0_text else "其他"

                    # Cell 2-6 不需要修改，因为它们是虚拟单元格

    def _apply_dedup_rules(
        self, template_config: Optional[TemplateMetadataModel], data_list: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """应用去重规则

        Args:
            template_config: 模板配置
            data_list: 排序后的数据列表

        Returns:
            去重后的数据列表
        """
        if not template_config or not template_config.dedup_rules:
            return data_list

        if not data_list:
            return data_list

        dedup_rules = template_config.dedup_rules
        # 提取所有需要去重的字段
        fields = [rule.field for rule in dedup_rules if rule.field]

        if not fields:
            return data_list

        seen = set()
        result = []

        for item in data_list:
            # 构建去重的 key（多字段组合）
            key_parts = []
            for field in fields:
                value = item.get(field, "")
                key_parts.append(str(value) if value is not None else "")
            key = tuple(key_parts)

            if key not in seen:
                seen.add(key)
                result.append(item)

        logger.info(f"Dedup rules: {len(data_list)} -> {len(result)} items, fields: {fields}")
        return result

    def _apply_product_matching(
        self, template_config: Optional[TemplateMetadataModel], data_list: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """应用产品匹配规则

        使用 product_matching 配置的标准产品列表进行匹配：
        - 匹配成功：标记 _seq，按 seq 顺序排序
        - 未匹配：追加到末尾，保持原顺序

        Args:
            template_config: 模板配置
            data_list: 去重后的数据列表

        Returns:
            匹配后的数据列表
        """
        if not template_config or not template_config.product_matching:
            return data_list

        if not data_list:
            return data_list

        product_matching = template_config.product_matching
        match_field = product_matching.match_field or "WLMS"

        # 构建产品名称到序号的映射
        product_map = {}
        for p in product_matching.products:
            product_map[p.name] = p.seq

        matched = []
        unmatched = []

        for item in data_list:
            item_value = item.get(match_field, "")
            # 查找匹配的产品
            if item_value in product_map:
                matched.append({**item, "_seq": product_map[item_value], "_matched": True})
            else:
                unmatched.append({**item, "_seq": 999, "_matched": False})

        # 按 seq 排序（匹配的在前，未匹配的在后）
        matched.sort(key=lambda x: x.get("_seq", 999))
        unmatched.sort(key=lambda x: x.get("_seq", 999))

        result = matched + unmatched

        logger.info(f"Product matching: {len(data_list)} -> {len(result)} items, matched: {len(matched)}, unmatched: {len(unmatched)}")

        return result

    def _calculate_discount_and_group(
        self,
        data_list: List[Dict[str, Any]],
        template_config: Optional[TemplateMetadataModel] = None
    ) -> tuple:
        """计算折扣分组

        支持两种模式：
        1. discount_templates 模式：按 seq_range 分组，使用配置的话术模板
        2. discount_template 模式（兼容）：按折扣率分组

        Args:
            data_list: 数据列表
            template_config: 模板配置

        Returns:
            (分组后的数据列表, 合并信息字典)
        """
        if not data_list:
            return data_list, {}

        # 优先使用 discount_templates（列表形式）
        if template_config and template_config.discount_templates:
            return self._group_by_seq_range(data_list, template_config.discount_templates)

        # 兼容旧的 discount_template（单个字符串形式）
        if template_config and template_config.discount_template:
            return self._group_by_discount_rate(data_list)

        return data_list, {}

    def _group_by_seq_range(
        self,
        data_list: List[Dict[str, Any]],
        discount_templates: List
    ) -> tuple:
        """按 seq_range 分组

        Args:
            data_list: 数据列表
            discount_templates: 折扣话术模板列表

        Returns:
            (分组后的数据列表, 合并信息字典)
        """
        # 构建 seq -> template 映射
        seq_to_template = {}
        for dt in discount_templates:
            for seq in dt.seq_range:
                seq_to_template[seq] = dt.template

        # 按 seq 排序
        sorted_data = sorted(data_list, key=lambda x: x.get("_seq", 999))

        # 计算合并信息
        merge_info = {}
        current_template = None
        span_count = 0
        start_idx = 0

        for idx, item in enumerate(sorted_data):
            seq = item.get("_seq", 999)
            template = seq_to_template.get(seq)

            if template != current_template:
                # 保存之前的合并信息
                if current_template is not None and span_count > 1:
                    merge_info[start_idx] = (current_template, span_count)
                # 开始新的分组
                current_template = template
                span_count = 1
                start_idx = idx
            else:
                span_count += 1

        # 保存最后一组的合并信息
        if current_template is not None and span_count > 1:
            merge_info[start_idx] = (current_template, span_count)

        logger.info(f"Sequence range grouping: {len(sorted_data)} items, {len(merge_info)} merge groups")
        return sorted_data, merge_info

    def _group_by_discount_rate(self, data_list: List[Dict[str, Any]]) -> tuple:
        """按折扣率分组（兼容旧逻辑）

        Args:
            data_list: 数据列表

        Returns:
            (分组后的数据列表, 合并信息字典)
        """
        # 为每条数据计算折扣率
        for item in data_list:
            jcjxj = item.get("JCJXJ")
            jczbj = item.get("JCZBJ")
            if jcjxj and jczbj and float(jczbj) > 0:
                discount = round(float(jcjxj) / float(jczbj) * 100)
                item["_discount"] = discount
            else:
                item["_discount"] = None

        # 按折扣率排序（None放在最后）
        sorted_data = sorted(data_list, key=lambda x: (x["_discount"] is None, x["_discount"] or 0))

        # 计算合并信息，并添加组内序号
        merge_info = {}
        current_discount = None
        span_count = 0
        start_idx = 0
        group_index = 0  # 组内序号

        for idx, item in enumerate(sorted_data):
            discount = item.get("_discount")
            if discount != current_discount:
                # 保存之前的合并信息
                if current_discount is not None and span_count > 1:
                    merge_info[start_idx] = (current_discount, span_count)
                # 开始新的分组
                current_discount = discount
                span_count = 1
                start_idx = idx
                group_index = 1  # 重置组内序号
            else:
                span_count += 1
                group_index += 1

            # 设置组内序号
            item["_group_index"] = group_index

        # 保存最后一组的合并信息
        if current_discount is not None and span_count > 1:
            merge_info[start_idx] = (current_discount, span_count)

        logger.info(f"Discount grouping: {len(sorted_data)} items, {len(merge_info)} merge groups")

        return sorted_data, merge_info

    def _apply_detail_filter(
        self, template_config: Optional[TemplateMetadataModel], data_list: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """应用明细过滤规则

        Args:
            template_config: 模板配置
            data_list: 原始数据列表

        Returns:
            过滤后的数据列表
        """
        logger.info(
            f"Applying detail filter, template_config: {template_config is not None}, detail_filter: {template_config.detail_filter if template_config else None}"
        )

        if not template_config or not template_config.detail_filter:
            logger.info("No detail filter configured, returning all data")
            return data_list

        filter_config = template_config.detail_filter

        # 使用 get_rules() 方法兼容 filter_rules 和 condition_groups
        filter_rules = filter_config.get_rules()
        if not filter_rules:
            return data_list

        # 处理条件组（组之间是 OR 关系）
        # 每个组内的条件是 AND 关系
        all_filtered = []

        for rule in filter_rules:
            # 组内条件（AND 关系）
            group_conditions = []
            for cond in rule.conditions:
                group_conditions.append(
                    FilterCondition(
                        field=cond.field,
                        operator=cond.operator,
                        value=cond.value,
                        value_type=cond.value_type,
                    )
                )

            if group_conditions:
                # 过滤符合该组条件的数据
                filtered = self.data_filler.filter_data(data_list, group_conditions)
                all_filtered.extend(filtered)

        # 去重
        seen = set()
        unique_filtered = []
        for item in all_filtered:
            key = item.get("WLDM") or item.get("ID") or id(item)
            if key not in seen:
                seen.add(key)
                unique_filtered.append(item)

        logger.info(f"Filtered {len(data_list)} rows to {len(unique_filtered)} rows")
        return unique_filtered

    def _apply_sort_rules(
        self,
        template_config: Optional[TemplateMetadataModel],
        data_list: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        """应用排序规则

        Args:
            template_config: 模板配置
            data_list: 原始数据列表

        Returns:
            排序后的数据列表
        """
        if not template_config or not template_config.sort_rules:
            return data_list

        sort_rules = template_config.sort_rules
        logger.info(f"Applying sort rules: {[r.name for r in sort_rules]}")

        def sort_key(item: Dict[str, Any]) -> tuple:
            keys = []
            for rule in sort_rules:
                field = rule.field
                order = rule.order.lower() if rule.order else "asc"
                value = item.get(field)
                # 将 None 转换为空字符串用于排序
                if value is None:
                    value = ""
                else:
                    value = str(value)
                # 升序用原值，降序用 (1, value) 使其排在后面
                if order == "desc":
                    keys.append((1, value))
                else:
                    keys.append((0, value))
            return tuple(keys)

        sorted_data = sorted(data_list, key=sort_key)
        return sorted_data

    def _parse_table_columns_from_config(self, table_config) -> List[Dict[str, str]]:
        """从模板配置解析表格列

        Args:
            table_config: 表格配置模型

        Returns:
            列配置列表
        """
        columns = []
        for col in table_config.columns:
            # auto_number 类型不需要 source_field
            field = col.source_field if col.type != ColumnType.AUTO_NUMBER else ""
            columns.append(
                {
                    "name": col.name,
                    "field": field,
                    "type": col.type,
                    "transform": col.transform,
                    "params": col.params,
                }
            )
        return columns

    def _parse_table_columns(self, table) -> List[Dict[str, str]]:
        """解析表格列配置"""
        columns = []

        if not table.rows:
            return columns

        # 第一行是表头
        header_row = table.rows[0]

        for idx, cell in enumerate(header_row.cells):
            text = cell.text.strip()
            # 移除占位符标记
            text = re.sub(r"\{\{.*\}\}", "", text).strip()

            # 推断字段
            field = self._infer_field(text, idx)
            col_type = "auto_number" if "序号" in text else "text"

            columns.append({"name": text, "field": field, "type": col_type})

        return columns

    def _infer_field(self, header_text: str, col_idx: int) -> str:
        """从表头推断字段

        使用 constants.HEADER_TO_FIELD
        """
        for key, field in HEADER_TO_FIELD.items():
            if key in header_text:
                return field

        return f"col_{col_idx}"

    def _replace_variables(self, text: str, data: Dict[str, Any]) -> str:
        """替换变量占位符

        使用 constants.DEFAULT_SPEECH_VARIABLES
        """
        for var_name, default_value in DEFAULT_SPEECH_VARIABLES.items():
            placeholder = f"{{{{{var_name}}}}}"
            if placeholder in text:
                value = data.get(var_name, default_value)
                text = text.replace(placeholder, str(value))

        return text

    def _fix_table_layout(self, doc: Document) -> None:
        """修复表格布局：将 autofit 改为 fixed，防止 WPS 渲染时列宽随内容变化

        Args:
            doc: Word 文档对象
        """
        for table in doc.tables:
            tbl = table._tbl
            tblPr = tbl.find(".//{*}tblPr")
            if tblPr is not None:
                tblLayout = tblPr.find(".//{*}tblLayout")
                if tblLayout is not None:
                    # 将 autofit 改为 fixed
                    tblLayout.set(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type",
                        "fixed",
                    )

    def _process_speeches(
        self,
        doc: Document,
        template_config: Optional[TemplateMetadataModel],
        detail_data_list: List[Dict[str, Any]],
    ) -> None:
        """处理话术占位符（完整集成）

        Args:
            doc: Word 文档对象
            template_config: 模板配置（可选）
            detail_data_list: 明细数据列表（用于话术条件判断）
        """
        # 如果没有模板配置，简化处理
        if not template_config or not template_config.speeches:
            self._process_speeches_simple(doc)
            return

        # 转换为 Speech 对象
        speeches = []
        for s in template_config.speeches:
            speech = Speech(
                id=s.id,
                type=s.type,
                content=s.content,
                mutex_group=s.mutex_group,
                variables={v.name: v.default for v in s.variables},
                conditions=[
                    FilterCondition(field=c.field, operator=c.operator, value=c.value)
                    for c in s.conditions
                ],
            )
            speeches.append(speech)

        # 处理话术 - 传递明细数据列表用于条件判断
        speech_contents = self.speech_processor.process_speeches(speeches, detail_data_list)

        # 填充话术占位符
        for para in doc.paragraphs:
            text = para.text

            if "{{#话术}}" in text or "{{话术}}" in text:
                # 合并所有话术内容
                full_speech = "\n".join(speech_contents)
                text = text.replace("{{#话术}}", "")
                text = text.replace("{{/话术}}", "")
                text = text.replace("{{话术}}", full_speech)

            if text != para.text:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = text

    def _process_speeches_simple(self, doc: Document) -> None:
        """简化处理话术占位符（无模板配置时）"""
        for para in doc.paragraphs:
            text = para.text

            if "{{#话术}}" in text or "{{话术}}" in text:
                # 简化：直接移除占位符
                text = text.replace("{{#话术}}", "")
                text = text.replace("{{/话术}}", "")
                # 移除话术变量占位符
                for var_name in DEFAULT_SPEECH_VARIABLES.keys():
                    placeholder = f"{{{{{var_name}}}}}"
                    text = text.replace(placeholder, DEFAULT_SPEECH_VARIABLES[var_name])

            if text != para.text:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = text

    def _process_table_speeches(self, table) -> None:
        """处理表格中的话术占位符"""
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if "{{#话术}}" in text or "{{话术}}" in text or "{{/话术}}" in text:
                        # 清除话术标记
                        text = text.replace("{{#话术}}", "")
                        text = text.replace("{{/话术}}", "")
                        # 替换话术变量
                        for var_name, default_value in DEFAULT_SPEECH_VARIABLES.items():
                            placeholder = f"{{{{{var_name}}}}}"
                            text = text.replace(placeholder, default_value)

                        # 更新段落文本
                        if text != para.text:
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = text

    def _get_speech_contents(
        self,
        template_config: Optional[TemplateMetadataModel],
        detail_data_list: List[Dict[str, Any]],
    ) -> List[str]:
        """获取话术内容列表

        Args:
            template_config: 模板配置（可选）
            detail_data_list: 明细数据列表（用于话术条件判断）

        Returns:
            话术内容列表
        """
        # 如果没有模板配置，返回空列表
        if not template_config or not template_config.speeches:
            return []

        # 转换为 Speech 对象
        speeches = []
        for s in template_config.speeches:
            speech = Speech(
                id=s.id,
                type=s.type,
                content=s.content,
                mutex_group=s.mutex_group,
                variables={v.name: v.default for v in s.variables},
                conditions=[
                    FilterCondition(field=c.field, operator=c.operator, value=c.value)
                    for c in s.conditions
                ],
            )
            speeches.append(speech)

        # 处理话术 - 传递明细数据列表用于条件判断
        speech_contents = self.speech_processor.process_speeches(speeches, detail_data_list)
        return speech_contents


class MultiDocumentGenerator:
    """多文档生成器 - 生成多个文档并打包"""

    def __init__(self, template_dir: str = "templates"):
        self.generator = DocumentGenerator(template_dir)

    def generate_batch(
        self,
        templates: List[TemplateRule],
        quote_data: Dict[str, Any],
        detail_data_list: List[Dict[str, Any]],
        output_dir: str = "output",
    ) -> List[GenerationResult]:
        """批量生成文档

        Args:
            templates: 模板列表
            quote_data: 报价单数据
            detail_data_list: 明细数据列表
            output_dir: 输出目录

        Returns:
            生成结果列表
        """
        results = []

        for template in templates:
            result = self.generator.generate(template, quote_data, detail_data_list, output_dir)
            results.append(result)

            if result.success:
                logger.info(f"Generated: {result.template_id} -> {result.file_path}")
            else:
                logger.error(f"Failed: {result.template_id} -> {result.error}")

        return results
